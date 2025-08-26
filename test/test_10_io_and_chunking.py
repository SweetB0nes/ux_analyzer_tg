# tests/test_10_io_and_chunking.py
from __future__ import annotations
from pathlib import Path
import pytest
from ux_analyzer_lc.utils.io import load_transcripts, load_brief_any

def test_transcripts_txt_docx(tmp_path: Path):
    (tmp_path / "t1.txt").write_text("Interviewer: hi\nUser: hello", encoding="utf-8")
    try:
        from docx import Document
        d = Document()
        d.add_paragraph("Интервьюер: привет")
        d.add_paragraph("Пользователь: ок")
        p = tmp_path / "t2.docx"; d.save(str(p))
    except Exception:
        pass
    trs = load_transcripts(str(tmp_path))
    assert len(trs) >= 1
    # длина первого транскрипта > 0
    assert len(trs[0][1]) > 0

def test_brief_merge_yaml_text(tmp_path: Path):
    (tmp_path / "b.yaml").write_text("project:{title:A,company:B}\ngoals:[g1]", encoding="utf-8")
    (tmp_path / "raw.txt").write_text("сырой бриф текст", encoding="utf-8")
    brief = load_brief_any([tmp_path/"b.yaml", tmp_path/"raw.txt"])
    assert brief["project"]["title"] == "A"
    assert "raw_brief" in brief
