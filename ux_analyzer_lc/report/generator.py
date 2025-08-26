from __future__ import annotations
from pathlib import Path
from jinja2 import Environment, FileSystemLoader, select_autoescape
from datetime import datetime
from weasyprint import HTML
from docx import Document
from ..config import SETTINGS

def _env():
    templates_dir = Path(__file__).parent / "templates"
    env = Environment(loader=FileSystemLoader(str(templates_dir)),
                      autoescape=select_autoescape(['html','xml']))
    return env

def render_html(payload: dict, company_name: str, author: str, report_title: str) -> str:
    tpl = _env().get_template("report.html.j2")
    return tpl.render(
        title=report_title,
        report_title=report_title,
        company_name=company_name,
        author=author,
        date=datetime.now().strftime("%Y-%m-%d"),
        **payload
    )

def save_all(payload: dict, company_name: str, author: str, report_title: str) -> dict:
    out_dir = Path(SETTINGS.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = out_dir / f"ux_report_{ts}"
    html_str = render_html(payload, company_name, author, report_title)
    html_path = base.with_suffix(".html")
    html_path.write_text(html_str, encoding="utf-8")

    # PDF
    pdf_path = base.with_suffix(".pdf")
    try:
        HTML(string=html_str).write_pdf(str(pdf_path))
    except Exception:
        pdf_path = None

    # DOCX (краткая версия)
    docx_path = base.with_suffix(".docx")
    try:
        doc = Document()
        doc.add_heading(report_title, 0)
        doc.add_paragraph(f"{company_name} • {author} • {datetime.now().strftime('%Y-%m-%d')}")
        doc.add_heading("Executive Summary", level=1)
        ff = payload.get("final", {}).get("final_findings", {})
        if isinstance(ff, dict):
            doc.add_paragraph(ff.get("executive_summary","—"))
        else:
            doc.add_paragraph("—")
        doc.save(str(docx_path))
    except Exception:
        docx_path = None

    return {"html": str(html_path), "pdf": str(pdf_path) if pdf_path else None, "docx": str(docx_path) if docx_path else None}
